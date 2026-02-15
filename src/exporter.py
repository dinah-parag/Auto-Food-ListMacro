def export_txt(df, path):
    with open(path, "w", encoding="utf-8") as f:
        for _, row in df.iterrows():
            f.write(f"{row['ingredient']} - {row['quantity']} {row['unit']}\n")


from docx import Document

def export_docx(df, path):
    doc = Document()
    doc.add_heading("Lista de Compras", level=1)

    for _, row in df.iterrows():
        doc.add_paragraph(f"{row['ingredient']} - {row['quantity']} {row['unit']}")

    doc.save(path)


from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

def export_pdf(df, path):
    doc = SimpleDocTemplate(path)
    elements = []
    styles = getSampleStyleSheet()

    for _, row in df.iterrows():
        elements.append(Paragraph(
            f"{row['ingredient']} - {row['quantity']} {row['unit']}",
            styles["Normal"]
        ))

    doc.build(elements)
